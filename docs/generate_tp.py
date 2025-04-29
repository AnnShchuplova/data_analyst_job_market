"""
Генератор DOCX «Текст программы» RU.17701729.11.03-01 12 01-1
ГОСТ 19.401-78, ГОСТ 19.105-78, ГОСТ 19.104-78, ГОСТ 2.104-2006
"""

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOC_CODE = "RU.17701729.11.03-01 12 01-1"
REPO_URL = "https://github.com/AnnShchuplova/data_analyst_job_market"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_page_margins(doc, top=20, bottom=20, left=30, right=15):
    """ГОСТ: поля 30-15-20-20 мм."""
    for section in doc.sections:
        section.top_margin    = Mm(top)
        section.bottom_margin = Mm(bottom)
        section.left_margin   = Mm(left)
        section.right_margin  = Mm(right)
        section.page_width    = Mm(210)
        section.page_height   = Mm(297)


def fmt(run, bold=False, size=12, font="Times New Roman", italic=False, underline=False):
    run.font.name       = font
    run.font.size       = Pt(size)
    run.font.bold       = bold
    run.font.italic     = italic
    run.font.underline  = underline
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    font)
    rFonts.set(qn("w:hAnsi"),    font)
    rFonts.set(qn("w:cs"),       font)
    rPr.insert(0, rFonts)


def para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=12,
         space_before=0, space_after=6, italic=False, underline=False,
         first_line=0, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line:
        p.paragraph_format.first_line_indent = Mm(first_line)
    if keep:
        p.paragraph_format.keep_with_next = True
    if text:
        r = p.add_run(text)
        fmt(r, bold=bold, size=size, italic=italic, underline=underline)
    return p


def heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 12}
    p = para(doc, text, bold=True, size=sizes.get(level, 12),
             align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=12, space_after=6, keep=True)
    return p


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        val = kwargs.get(edge, {})
        if val:
            e = OxmlElement(f"w:{edge}")
            for k, v in val.items():
                e.set(qn(f"w:{k}"), str(v))
            tcBorders.append(e)
    tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.alignment    = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if p.runs:
        p.runs[0].text = text
        fmt(p.runs[0], bold=bold, size=size)
    else:
        r = p.add_run(text)
        fmt(r, bold=bold, size=size)


def add_table(doc, headers, rows, col_widths=None):
    """Таблица с заголовком (шапкой) и данными."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style     = "Table Grid"

    border = {"val": "single", "sz": "4", "color": "000000"}

    # Шапка
    hdr = table.rows[0]
    hdr.height = Mm(10)
    for i, h in enumerate(headers):
        cell_text(hdr.cells[i], h, bold=True, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_border(hdr.cells[i],
                        top=border, left=border, bottom=border, right=border)

    # Данные
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell_text(row.cells[c_idx], val, size=10)
            set_cell_border(row.cells[c_idx],
                            top=border, left=border,
                            bottom=border, right=border)

    # Ширины колонок
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Mm(w)
    return table


def page_break(doc):
    doc.add_page_break()


def stamp_para(doc):
    """
    Нижний штамп (упрощённый текстовый вариант).
    В реальном документе штамп вставляется как таблица в нижний колонтитул;
    здесь — строка в тексте для наглядности.
    """
    p = para(doc, DOC_CODE,
             align=WD_ALIGN_PARAGRAPH.LEFT,
             size=9, space_before=0, space_after=0)
    return p


def add_footer_stamp(doc):
    """Добавляет штамп-таблицу в нижний колонтитул каждой секции."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Очищаем стандартный параграф колонтитула
        for p in footer.paragraphs:
            p.clear()

        ft = footer.add_table(rows=2, cols=5, width=Mm(165))
        ft.alignment = WD_TABLE_ALIGNMENT.LEFT
        ft.style     = "Table Grid"

        border = {"val": "single", "sz": "4", "color": "000000"}

        # Строка 1: Изм. | Лист | № докум. | Подп. | Дата
        labels1 = ["Изм.", "Лист", "№ докум.", "Подп.", "Дата"]
        widths1 = [15, 15, 40, 20, 25]
        for i, (lbl, w) in enumerate(zip(labels1, widths1)):
            c = ft.rows[0].cells[i]
            c.width = Mm(w)
            cell_text(c, lbl, bold=False, size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_border(c, top=border, left=border,
                            bottom=border, right=border)

        # Строка 2: код документа + пустые ячейки
        ft.rows[1].cells[0].merge(ft.rows[1].cells[4])
        c = ft.rows[1].cells[0]
        cell_text(c, DOC_CODE, bold=False, size=9,
                  align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_border(c, top=border, left=border,
                        bottom=border, right=border)


# ---------------------------------------------------------------------------
# Страницы документа
# ---------------------------------------------------------------------------

def add_title_page(doc):
    """Титульный лист (ЛУ) по ГОСТ 19.104-78."""
    para(doc, "ПРАВИТЕЛЬСТВО РОССИЙСКОЙ ФЕДЕРАЦИИ",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    para(doc, "НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    para(doc, "«ВЫСШАЯ ШКОЛА ЭКОНОМИКИ»",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=6)
    para(doc, "Факультет компьютерных наук",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "Департамент программной инженерии",
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # Согласовано / Утверждаю
    t = doc.add_table(rows=4, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [
        ["СОГЛАСОВАНО", "УТВЕРЖДАЮ"],
        ["Доцент департамента программной инженерии\nфакультета компьютерных наук",
         "Академический руководитель\nобразовательной программы\n«Программная инженерия»"],
        ["____________ А.В. Меликян", "____________ Н.А. Павлочев"],
        ["«__» _____________ 2026 г.", "«__» _____________ 2026 г."],
    ]
    widths = [80, 80]
    for r_idx, row_data in enumerate(labels):
        for c_idx, txt in enumerate(row_data):
            c = t.rows[r_idx].cells[c_idx]
            c.width = Mm(widths[c_idx])
            cell_text(c, txt, bold=(r_idx == 0), size=12,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "", space_before=18, space_after=18)

    para(doc, "ПРОГРАММА ДЛЯ АНАЛИЗА РЫНКА ТРУДА АНАЛИТИКОВ ДАННЫХ",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=6)
    para(doc, "Текст программы",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=6)
    para(doc, "ЛИСТ УТВЕРЖДЕНИЯ",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=12)
    para(doc, DOC_CODE + "-ЛУ",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=48)

    # Исполнитель (правый край)
    p = para(doc, "", space_after=0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Исполнитель\nСтудент группы БПИ247\n____________ Н.А. Тищенко\n«__» _____________ 2026 г.")
    fmt(r, size=12)

    para(doc, "", space_before=48, space_after=0)
    para(doc, "Москва 2026",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=12)
    page_break(doc)


def add_lu_back(doc):
    """Оборот ЛУ."""
    para(doc, "УТВЕРЖДЁН", bold=False, size=12, space_after=0)
    para(doc, DOC_CODE + "-ЛУ", bold=False, size=12, space_after=36)
    para(doc, "ПРОГРАММА ДЛЯ АНАЛИЗА РЫНКА ТРУДА АНАЛИТИКОВ ДАННЫХ",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=6)
    para(doc, "Текст программы",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=6)
    para(doc, DOC_CODE,
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=12)
    para(doc, "Листов 9",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=48)
    para(doc, "Москва 2026",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    page_break(doc)


def add_annotation(doc):
    heading(doc, "АННОТАЦИЯ", 1)
    text = (
        "Настоящий документ является программным документом вида «Текст программы» "
        "и разработан в соответствии с требованиями ГОСТ 19.401-78.\n\n"
        "Документ распространяется на программу «Программа для анализа рынка труда "
        "аналитиков данных» (далее — DataTrack), разработанную студентом группы "
        "БПИ247 НИУ ВШЭ Тищенко Н.А.\n\n"
        "В соответствии с п. 1.2 ГОСТ 19.401-78 допускается не включать в документ "
        "текст программы, если он полностью совпадает с текстом, переданным на хранение "
        "в репозиторий программной документации. В данном документе вместо листинга "
        "исходных текстов приводится ссылка на общедоступный репозиторий системы "
        "контроля версий, где хранятся актуальные исходные тексты всех компонентов "
        "программы.\n\n"
        "Документ содержит сведения о составе, структуре, языке реализации и "
        "местонахождении всех исходных программных модулей, необходимые для их "
        "однозначной идентификации и воспроизведения рабочего окружения."
    )
    for block in text.split("\n\n"):
        para(doc, block, space_before=0, space_after=6, first_line=12.5)
    page_break(doc)


def add_toc(doc):
    heading(doc, "СОДЕРЖАНИЕ", 1)
    entries = [
        ("АННОТАЦИЯ", "3"),
        ("1. ОБЩИЕ СВЕДЕНИЯ", "4"),
        ("   1.1. Назначение документа", "4"),
        ("   1.2. Местонахождение текстов программы", "4"),
        ("   1.3. Язык программирования", "4"),
        ("   1.4. Операционная среда", "5"),
        ("2. СОСТАВ ПРОГРАММЫ", "5"),
        ("   2.1. Точка входа", "5"),
        ("   2.2. Модуль веб-интерфейса (app/)", "5"),
        ("   2.3. Модули предметной области и инфраструктуры (src/)", "6"),
        ("   2.4. Скрипты обслуживания (scripts/)", "7"),
        ("   2.5. Тесты (tests/)", "8"),
        ("   2.6. Вспомогательные файлы", "8"),
        ("3. СВЕДЕНИЯ О ЗАВИСИМОСТЯХ", "9"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Mm(155), WD_ALIGN_PARAGRAPH.RIGHT)
        r = p.add_run(title + "\t" + page)
        fmt(r, size=12)
    page_break(doc)


def add_section1(doc):
    heading(doc, "1. ОБЩИЕ СВЕДЕНИЯ", 1)

    heading(doc, "1.1. Назначение документа", 2)
    para(doc,
         "Документ устанавливает состав исходных программных модулей, язык их "
         "реализации и местонахождение в системе контроля версий.",
         first_line=12.5, space_after=6)

    heading(doc, "1.2. Местонахождение текстов программы", 2)
    para(doc,
         "Полные актуальные тексты всех программных модулей, входящих в состав "
         "системы DataTrack, размещены в общедоступном репозитории системы контроля "
         "версий Git по следующему адресу:",
         first_line=12.5, space_after=6)

    p = para(doc, REPO_URL,
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12,
             space_before=6, space_after=6)
    p.runs[0].font.underline = True
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    para(doc, "Доступ: открытый (public repository).", first_line=12.5, space_after=3)
    para(doc, "Платформа: GitHub (Microsoft Corporation).", first_line=12.5, space_after=3)
    para(doc, "Основная ветка: main.", first_line=12.5, space_after=6)
    para(doc,
         "Для получения локальной копии исходных текстов необходимо выполнить команду:",
         first_line=12.5, space_after=3)

    p = para(doc,
             "    git clone " + REPO_URL + ".git",
             size=11, space_before=3, space_after=6)
    p.runs[0].font.name = "Courier New"

    para(doc,
         "Фиксация (commit), соответствующая версии настоящего документа, "
         "идентифицируется по хеш-сумме последнего коммита ветки main на дату "
         "утверждения документа.",
         first_line=12.5, space_after=6)

    heading(doc, "1.3. Язык программирования", 2)
    para(doc, "Все программные модули системы написаны на языке Python версии 3.11 и выше.",
         first_line=12.5, space_after=3)
    para(doc, "Стандарт кодирования исходных файлов: UTF-8 (без BOM).",
         first_line=12.5, space_after=3)
    para(doc, "Стиль оформления кода: PEP 8.",
         first_line=12.5, space_after=6)

    heading(doc, "1.4. Операционная среда", 2)
    para(doc,
         "Программа функционирует под управлением ОС Linux (в контейнерной среде Docker) "
         "и Windows 10/11 при локальном запуске. Совместимость обеспечивается через "
         "виртуальное окружение Python (venv / pip).",
         first_line=12.5, space_after=6)
    page_break(doc)


def add_section2(doc):
    heading(doc, "2. СОСТАВ ПРОГРАММЫ", 1)

    heading(doc, "2.1. Точка входа", 2)
    add_table(doc,
        ["Файл", "Назначение"],
        [
            ["run.py",
             "Точка запуска Streamlit-приложения; обеспечивает совместимость с PyInstaller."],
            ["app/main.py",
             "Инициализация страницы Streamlit, маршрутизация между представлениями "
             "через st.session_state['page'], рендеринг навигационной панели (render_top_nav)."],
        ],
        col_widths=[55, 110]
    )
    para(doc, "", space_after=6)

    heading(doc, "2.2. Модуль веб-интерфейса (app/)", 2)
    add_table(doc,
        ["Файл", "Назначение"],
        [
            ["app/views/home.py",
             "Представление главной страницы: сводная статистика по датасету (view_home)."],
            ["app/views/aggregation.py",
             "Представление агрегации: 6 вкладок фильтрации, вкладка «Мой профиль», "
             "экспорт CSV (view_aggregation)."],
            ["app/views/clusters.py",
             "Представление кластеризации: выбор признаков, запуск ClusteringService, "
             "отображение карточек кластеров (view_clusters)."],
            ["app/views/arima.py",
             "Представление временных трендов: SARIMA-прогноз, доверительные интервалы "
             "(view_arima)."],
            ["app/views/predict_salary.py",
             "Представление прогноза зарплаты: обучение SalaryPredictor на лету, "
             "предсказание по профилю (view_salary_predictor)."],
            ["app/utils/visualizations.py",
             "Вспомогательные функции построения Plotly-графиков."],
            ["app/utils/data_loader.py",
             "Утилиты загрузки данных для слоя представлений."],
        ],
        col_widths=[55, 110]
    )
    para(doc, "", space_after=6)
    page_break(doc)

    heading(doc, "2.3. Модули предметной области и инфраструктуры (src/)", 2)

    heading(doc, "2.3.1. Доменные модели (src/domain/)", 3)
    add_table(doc,
        ["Файл", "Назначение"],
        [["src/domain/models.py",
          "Датаклассы системы: HomeStats, ClusterEntity, ClusteringResult, "
          "SalaryPredictionResult."]],
        col_widths=[55, 110]
    )
    para(doc, "", space_after=6)

    heading(doc, "2.3.2. Инфраструктура (src/infrastructure/)", 3)
    add_table(doc,
        ["Файл", "Назначение"],
        [
            ["src/infrastructure/vacancies_repository.py",
             "Доступ к данным: загрузка последнего CSV из finaldata/, "
             "вычисление MD5-checksum для инвалидации кэша."],
            ["src/infrastructure/cache.py",
             "Кэш результатов кластеризации (ClusteringCache, joblib)."],
        ],
        col_widths=[65, 100]
    )
    para(doc, "", space_after=6)

    heading(doc, "2.3.3. Сервисный слой (src/services/)", 3)
    add_table(doc,
        ["Файл", "Назначение"],
        [
            ["src/services/home_service.py",
             "Вычисление агрегированной статистики главной страницы."],
            ["src/services/clustering_service.py",
             "Оркестрация кластеризации: автовыбор K-Means / K-Modes / K-Prototypes, "
             "оптимизация K по Silhouette Score."],
            ["src/services/mock.py",
             "Заглушка сервисного слоя (legacy, не используется в рабочих путях)."],
        ],
        col_widths=[65, 100]
    )
    para(doc, "", space_after=6)

    heading(doc, "2.3.4. Сбор данных (src/data_collection/)", 3)
    add_table(doc,
        ["Файл", "Назначение"],
        [
            ["src/data_collection/hh_parser.py",
             "HTTP-клиент hh.ru API: постраничный обход (до 20 страниц × 100 вакансий), "
             "retry / exponential backoff, фильтрация по ID профессиональных ролей."],
            ["src/data_collection/filters.py",
             "Фильтрация вакансий по профессиональным ролям аналитиков данных, "
             "удаление дубликатов по полю id."],
        ],
        col_widths=[65, 100]
    )
    para(doc, "", space_after=6)

    heading(doc, "2.3.5. Обработка данных (src/data_processing/)", 3)
    add_table(doc,
        ["Файл", "Назначение"],
        [
            ["src/data_processing/cleaner.py",
             "Преобразование JSON → DataFrame: нормализация зарплат, опыта, навыков; "
             "удаление выбросов (IQR + абсолютный порог 15 000 ₽). "
             "Создаёт поля min_experience_years, avg_experience_years."],
            ["src/data_processing/feature_engineering.py",
             "Генерация ML-признаков: TF-IDF навыков и текстов требований, "
             "target encoding (только по train-выборке), бинарные флаги, "
             "теги грейда (is_senior / is_middle / is_junior), "
             "поле estimated_experience_years."],
        ],
        col_widths=[65, 100]
    )
    para(doc, "", space_after=6)
    page_break(doc)

    heading(doc, "2.3.6. ML-модели (src/ml/)", 3)
    add_table(doc,
        ["Файл", "Назначение"],
        [
            ["src/ml/salary_predictor.py",
             "Ансамбль RandomForest + CatBoost + GradientBoosting; взвешивание "
             "по R²; сериализация в models/salary_predictor.pkl (pickle)."],
            ["src/ml/arima_analyzer.py",
             "SARIMA-модель временного ряда числа вакансий; log1p-преобразование; "
             "Monte Carlo доверительные интервалы (1000 путей); "
             "сериализация в models/arima_model.joblib."],
        ],
        col_widths=[65, 100]
    )
    para(doc, "", space_after=12)

    heading(doc, "2.4. Скрипты обслуживания (scripts/)", 2)
    add_table(doc,
        ["Файл", "Назначение"],
        [
            ["scripts/run_pipeline.py",
             "Сквозной запуск: сбор → фильтрация → очистка → сохранение CSV."],
            ["scripts/collect_data.py",
             "Запуск HHParser, сохранение сырых JSON в data/raw/."],
            ["scripts/filter_data.py",
             "Фильтрация сырых JSON по ролям аналитиков данных."],
            ["scripts/train_models.py",
             "Обучение SalaryPredictor; запуск с флагом --compare выводит "
             "сравнительную таблицу метрик RF / CatBoost / GB / Ансамбль."],
            ["scripts/merde_databases.py",
             "Объединение нескольких CSV из finaldata/ в единый датасет. "
             "Примечание: в имени файла допущена опечатка "
             "(корректное имя — merge_databases.py)."],
            ["scripts/check_roles.py",
             "Утилита диагностики: вывод доступных ID профессиональных ролей hh.ru."],
        ],
        col_widths=[65, 100]
    )
    para(doc, "", space_after=12)

    heading(doc, "2.5. Тесты (tests/)", 2)
    add_table(doc,
        ["Файл", "Назначение"],
        [["tests/test_modules.py",
          "Модульные тесты (unittest / pytest): TestHHParser, TestFilters, TestCleaner. "
          "Запуск: python -m pytest tests/test_modules.py"]],
        col_widths=[65, 100]
    )
    para(doc, "", space_after=12)
    page_break(doc)

    heading(doc, "2.6. Вспомогательные файлы", 2)
    add_table(doc,
        ["Файл", "Назначение"],
        [
            ["docker-compose.yml",
             "Описание двух сервисов Docker: web (порт 8501:8501) и parser (без портов)."],
            ["docker/web.Dockerfile",   "Образ Streamlit-приложения."],
            ["docker/parser.Dockerfile","Образ контейнера парсера данных."],
            ["requirements.txt",
             "Список зависимостей pip. Примечание: библиотека catboost указана дважды "
             "(строки 5 и 19); дублирование подлежит устранению."],
            [".env",
             "Переменные окружения: HH_USER_AGENT, LOG_LEVEL. "
             "Файл не хранится в репозитории (добавлен в .gitignore)."],
        ],
        col_widths=[55, 110]
    )
    para(doc, "", space_after=6)
    page_break(doc)


def add_section3(doc):
    heading(doc, "3. СВЕДЕНИЯ О ЗАВИСИМОСТЯХ", 1)
    para(doc,
         "Перечень внешних программных библиотек задаётся файлом requirements.txt "
         "в корне репозитория. Актуальный файл доступен по адресу:",
         first_line=12.5, space_after=6)

    p = para(doc,
             REPO_URL + "/blob/main/requirements.txt",
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=11,
             space_before=3, space_after=6)
    p.runs[0].font.underline = True
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    p.runs[0].font.name = "Courier New"

    para(doc,
         "Основные зависимости на момент утверждения документа приведены в таблице ниже.",
         first_line=12.5, space_after=6)

    add_table(doc,
        ["Библиотека", "Версия (min)", "Назначение"],
        [
            ["streamlit",       "—",        "Веб-интерфейс (фреймворк дашборда)"],
            ["pandas",          "≥ 2.2.3",  "Работа с табличными данными"],
            ["numpy",           "≥ 1.24.0", "Числовые вычисления"],
            ["scikit-learn",    "≥ 1.3.0",  "ML-алгоритмы, K-Means, метрики"],
            ["catboost",        "≥ 1.2.0",  "Градиентный бустинг (CatBoostRegressor)"],
            ["statsmodels",     "≥ 0.14.0", "SARIMA, ADF-тест"],
            ["kmodes",          "≥ 0.12.2", "K-Modes, K-Prototypes"],
            ["gower",           "≥ 0.1.2",  "Расстояние Говера для смешанных признаков"],
            ["plotly",          "≥ 5.17.0", "Интерактивные графики"],
            ["requests",        "≥ 2.31.0", "HTTP-клиент для hh.ru API"],
            ["joblib",          "≥ 1.3.0",  "Сериализация моделей (ARIMA, кэш)"],
            ["scipy",           "≥ 1.11.0", "Статистические функции"],
            ["pytest",          "≥ 7.4.0",  "Запуск модульных тестов"],
            ["python-dotenv",   "≥ 1.0.0",  "Загрузка переменных окружения из .env"],
            ["matplotlib",      "≥ 3.7.0",  "Статические графики"],
            ["seaborn",         "≥ 0.12.2", "Статистическая визуализация"],
        ],
        col_widths=[45, 35, 85]
    )
    para(doc, "", space_after=12)

    heading(doc, "Порядок установки зависимостей", 3)
    p = para(doc, "    pip install -r requirements.txt",
             size=11, space_before=3, space_after=6)
    p.runs[0].font.name = "Courier New"


def add_change_log(doc):
    page_break(doc)
    heading(doc, "ЛИСТ РЕГИСТРАЦИИ ИЗМЕНЕНИЙ", 1)
    add_table(doc,
        ["Изм.", "Номера листов (страниц)\nизменённых", "Номера листов (страниц)\nзамённых",
         "Всего листов", "№ документа", "Входящий № сопроводительного документа", "Дата"],
        [["", "", "", "", "", "", ""]],
        col_widths=[12, 30, 30, 20, 25, 28, 20]
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    set_page_margins(doc)

    # Стиль Normal по умолчанию
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_title_page(doc)
    add_lu_back(doc)
    add_toc(doc)
    add_annotation(doc)
    add_section1(doc)
    add_section2(doc)
    add_section3(doc)
    add_change_log(doc)
    add_footer_stamp(doc)

    out = "docs/ТП_RU.17701729.11.03-01_12_01-1.docx"
    doc.save(out)
    print(f"Сохранён: {out}")


if __name__ == "__main__":
    build()
